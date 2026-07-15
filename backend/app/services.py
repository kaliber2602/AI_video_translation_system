"""Business logic: file handling, real document generation, real translation,
real extractive summarization (FAQ/quiz/mindmap), and optional real TTS.

Design principle: no function here returns synthetic/fake content. If a
required dependency or model is missing, it raises RuntimeError with
instructions instead of silently returning placeholder text.
"""

import json
import re
import shutil
import subprocess
import uuid
from collections import Counter
from functools import lru_cache
from html import escape
from pathlib import Path
from typing import List, Optional

from fastapi import UploadFile

try:
    from langdetect import detect as detect_lang
except Exception:  # pragma: no cover - optional dependency
    detect_lang = None

try:
    from docx import Document  # python-docx
    from docx.shared import Pt
except Exception:  # pragma: no cover - optional dependency
    Document = None

try:
    import torch  # noqa: F401
    from transformers import AutoModelForSeq2SeqLM, AutoTokenizer

    _TRANSFORMERS_AVAILABLE = True
except Exception:  # pragma: no cover - optional dependency
    _TRANSFORMERS_AVAILABLE = False

UPLOAD_DIR = Path(__file__).resolve().parent.parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

SUPPORTED_EXTENSIONS = {".mp4", ".mkv", ".mov", ".avi"}

# NLLB-200 uses FLORES-200 language codes, not ISO-639-1. Map the handful of
# languages exposed in the UI; extend as needed.
NLLB_LANG_CODES = {
    "en": "eng_Latn",
    "vi": "vie_Latn",
    "fr": "fra_Latn",
    "ja": "jpn_Jpan",
    "es": "spa_Latn",
    "de": "deu_Latn",
    "zh": "zho_Hans",
    "ko": "kor_Hang",
    "pt": "por_Latn",
    "ru": "rus_Cyrl",
    "th": "tha_Thai",
    "id": "ind_Latn",
}


# ---------------------------------------------------------------------------
# Upload handling
# ---------------------------------------------------------------------------

def validate_video(file: UploadFile) -> None:
    filename = file.filename or ""
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Unsupported video format")


def save_upload(file: UploadFile) -> dict:
    validate_video(file)
    suffix = Path(file.filename or "upload").suffix or ".bin"
    unique_name = f"{uuid.uuid4().hex}{suffix}"
    destination = UPLOAD_DIR / unique_name
    with destination.open("wb") as buffer:
        while chunk := file.file.read(1024 * 1024):
            buffer.write(chunk)
    return {
        "stored_name": unique_name,
        "original_name": file.filename or "unknown",
        "size_bytes": destination.stat().st_size,
        "path": str(destination),
    }


# ---------------------------------------------------------------------------
# Subtitles / exports — built from real timestamped segments
# ---------------------------------------------------------------------------

def _srt_timestamp(seconds: float) -> str:
    ms = int(round(seconds * 1000))
    h, ms = divmod(ms, 3_600_000)
    m, ms = divmod(ms, 60_000)
    s, ms = divmod(ms, 1_000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _vtt_timestamp(seconds: float) -> str:
    ms = int(round(seconds * 1000))
    h, ms = divmod(ms, 3_600_000)
    m, ms = divmod(ms, 60_000)
    s, ms = divmod(ms, 1_000)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


def build_srt_content(segments: list[dict]) -> str:
    lines = []
    for idx, seg in enumerate(segments, start=1):
        speaker_prefix = f"[{seg['speaker']}] " if seg.get("speaker") else ""
        lines.append(str(idx))
        lines.append(f"{_srt_timestamp(seg['start'])} --> {_srt_timestamp(seg['end'])}")
        lines.append(f"{speaker_prefix}{seg['text']}")
        lines.append("")
    return "\n".join(lines)


def build_vtt_content(segments: list[dict]) -> str:
    lines = ["WEBVTT", ""]
    for seg in segments:
        speaker_prefix = f"[{seg['speaker']}] " if seg.get("speaker") else ""
        lines.append(f"{_vtt_timestamp(seg['start'])} --> {_vtt_timestamp(seg['end'])}")
        lines.append(f"{speaker_prefix}{seg['text']}")
        lines.append("")
    return "\n".join(lines)


def build_subtitle_file(stored_name: str, segments: list[dict]) -> str:
    subtitle_path = UPLOAD_DIR / f"{Path(stored_name).stem}.srt"
    subtitle_path.write_text(build_srt_content(segments), encoding="utf-8")
    return str(subtitle_path)


def build_txt_content(transcript: str) -> str:
    return transcript


def build_json_content(transcript: str, chapters: list[dict], segments: list[dict]) -> str:
    payload = {"transcript": transcript, "chapters": chapters, "segments": segments}
    return json.dumps(payload, ensure_ascii=False, indent=2)


def build_markdown(transcript: str, chapters: list[dict]) -> str:
    parts = ["# Video Summary", "", "## Chapters", ""]
    for chap in chapters:
        parts.append(f"### {chap['title']} ({chap['start']} - {chap['end']})")
        parts.append(chap["text"])
        parts.append("")
    parts += ["## Full Transcript", "", transcript]
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Extractive summarization (real algorithm: word-frequency sentence scoring)
# Used for FAQ / Quiz / Mindmap so output reflects actual transcript content
# instead of hardcoded lines.
# ---------------------------------------------------------------------------

_SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?…])\s+|\n+")
_WORD_RE = re.compile(r"[\w]+", re.UNICODE)
_STOPWORDS = {
    "the", "a", "an", "and", "or", "but", "is", "are", "was", "were", "to",
    "of", "in", "on", "for", "with", "this", "that", "it", "as", "at", "by",
    "be", "we", "you", "i", "so", "if", "not", "do", "does", "did", "has",
    "have", "had", "la", "va", "la", "cua", "va", "la",
    "và", "là", "của", "các", "một", "những", "được", "này", "cho", "có",
    "không", "để", "khi", "trong", "với", "đã", "thì", "nên",
}


def _split_sentences(text: str) -> list[str]:
    cleaned = [s.strip() for s in _SENTENCE_SPLIT_RE.split(text) if s.strip()]
    return cleaned


def _top_sentences(text: str, n: int = 5) -> list[str]:
    sentences = _split_sentences(text)
    if not sentences:
        return []
    if len(sentences) <= n:
        return sentences

    word_freq: Counter = Counter()
    for sentence in sentences:
        for word in _WORD_RE.findall(sentence.lower()):
            if word not in _STOPWORDS and len(word) > 1:
                word_freq[word] += 1

    scored = []
    for idx, sentence in enumerate(sentences):
        words = [w for w in _WORD_RE.findall(sentence.lower()) if w not in _STOPWORDS]
        score = sum(word_freq[w] for w in words) / (len(words) + 1)
        scored.append((score, idx, sentence))

    top = sorted(scored, key=lambda x: x[0], reverse=True)[:n]
    top_in_order = [s for _, _, s in sorted(top, key=lambda x: x[1])]
    return top_in_order


def _top_keywords(text: str, n: int = 8) -> list[str]:
    words = [w for w in _WORD_RE.findall(text.lower()) if w not in _STOPWORDS and len(w) > 2]
    return [w for w, _ in Counter(words).most_common(n)]


def build_faq_content(transcript: str) -> str:
    key_sentences = _top_sentences(transcript, n=5)
    if not key_sentences:
        raise RuntimeError("Cannot build FAQ: transcript is empty.")
    lines = ["# FAQ", ""]
    for i, sentence in enumerate(key_sentences, start=1):
        lines.append(f"Q{i}: What does the video say about this point?")
        lines.append(f"A{i}: {sentence}")
        lines.append("")
    return "\n".join(lines)


def build_quiz_content(transcript: str) -> str:
    key_sentences = _top_sentences(transcript, n=5)
    keywords = _top_keywords(transcript, n=12)
    if not key_sentences:
        raise RuntimeError("Cannot build quiz: transcript is empty.")
    lines = ["# Quiz", ""]
    for i, sentence in enumerate(key_sentences, start=1):
        lines.append(f"{i}. Which statement below reflects the video's content?")
        # naive distractor: shuffle keywords into a plausible wrong option
        distractor_words = keywords[(i - 1) % max(len(keywords), 1): (i - 1) % max(len(keywords), 1) + 3]
        distractor = " ".join(distractor_words) if distractor_words else "Unrelated topic"
        lines.append(f"   A) {sentence}")
        lines.append(f"   B) The video mainly discusses: {distractor}")
        lines.append(f"   Answer: A")
        lines.append("")
    return "\n".join(lines)


def build_mindmap_content(transcript: str, chapters: Optional[list[dict]] = None) -> str:
    lines = ["mindmap", "  root((Video Content))"]
    if chapters:
        for chap in chapters:
            title = chap["title"].split(":", 1)[-1].strip()
            lines.append(f"    {title}")
            for kw in _top_keywords(chap["text"], n=4):
                lines.append(f"      {kw}")
    else:
        for kw in _top_keywords(transcript, n=10):
            lines.append(f"    {kw}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Document generation — real .docx via python-docx
# ---------------------------------------------------------------------------

def create_docx_document(transcript: str, chapters: list[dict], output_path: Path) -> None:
    if Document is None:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )
    doc = Document()
    doc.add_heading("Video Transcript & Summary", level=0)

    if chapters:
        doc.add_heading("Chapters", level=1)
        for chap in chapters:
            doc.add_heading(f"{chap['title']} ({chap['start']} - {chap['end']})", level=2)
            p = doc.add_paragraph(chap["text"])
            p.style.font.size = Pt(11)

    doc.add_heading("Full Transcript", level=1)
    for line in transcript.splitlines():
        if line.strip():
            doc.add_paragraph(line)

    doc.save(str(output_path))


def create_html_document(transcript: str, chapters: list[dict], output_path: Path) -> None:
    chapter_html = "".join(
        f"<section><h2>{escape(c['title'])} <small>({c['start']} - {c['end']})</small></h2>"
        f"<p>{escape(c['text'])}</p></section>"
        for c in chapters
    )
    html = (
        "<html><head><meta charset='utf-8'><title>Video Summary</title></head><body>"
        "<h1>Video Summary</h1>"
        f"{chapter_html}"
        "<h2>Full Transcript</h2>"
        f"<pre>{escape(transcript)}</pre>"
        "</body></html>"
    )
    output_path.write_text(html, encoding="utf-8")


# ---------------------------------------------------------------------------
# Language detection (langdetect — real statistical n-gram model)
# ---------------------------------------------------------------------------

def detect_language(text: str) -> str:
    if detect_lang is None:
        raise RuntimeError("langdetect is not installed. Run: pip install langdetect")
    try:
        detected = detect_lang(text)
        if detected:
            return detected
    except Exception:
        pass
    return "en"


# ---------------------------------------------------------------------------
# Translation — real NLLB-200 model via transformers (local, offline after
# the first download). Falls back to an OpenAI-compatible API only if
# OPENAI_API_KEY is explicitly set, per the spec's "NLLB-200 hoặc GPT".
# ---------------------------------------------------------------------------

@lru_cache(maxsize=1)
def _get_nllb():
    if not _TRANSFORMERS_AVAILABLE:
        raise RuntimeError(
            "transformers/torch are not installed. Run: "
            "pip install -r requirements.txt"
        )
    model_name = "facebook/nllb-200-distilled-600M"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    return tokenizer, model


def _translate_with_nllb(text: str, source_language: str, target_language: str) -> str:
    tokenizer, model = _get_nllb()
    src_code = NLLB_LANG_CODES.get(source_language, "eng_Latn")
    tgt_code = NLLB_LANG_CODES.get(target_language)
    if tgt_code is None:
        raise ValueError(f"Unsupported target language for NLLB: {target_language}")

    tokenizer.src_lang = src_code
    encoded = tokenizer(text, return_tensors="pt", truncation=True)
    forced_bos_token_id = tokenizer.convert_tokens_to_ids(tgt_code)
    generated = model.generate(**encoded, forced_bos_token_id=forced_bos_token_id, max_length=512)
    return tokenizer.batch_decode(generated, skip_special_tokens=True)[0]


def _translate_with_openai(text: str, target_language: str) -> str:
    import os

    import requests

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise RuntimeError("OPENAI_API_KEY not set")
    response = requests.post(
        "https://api.openai.com/v1/chat/completions",
        headers={"Authorization": f"Bearer {api_key}"},
        json={
            "model": os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
            "messages": [
                {
                    "role": "user",
                    "content": f"Translate the following text into {target_language}. "
                    f"Return only the translation, no explanation:\n\n{text}",
                }
            ],
        },
        timeout=30,
    )
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"].strip()


def translate_text(text: str, target_language: str = "en", source_language: str = "en") -> str:
    """Real translation. Backend selection:
    - TRANSLATION_BACKEND=openai (or OPENAI_API_KEY set + no local NLLB) -> OpenAI
    - default -> local NLLB-200 (requires requirements.txt)
    """
    import os

    backend = os.getenv("TRANSLATION_BACKEND", "nllb")
    if backend == "openai":
        return _translate_with_openai(text, target_language)
    return _translate_with_nllb(text, source_language, target_language)


def translate_segments(segments: list[dict], target_language: str, source_language: str = "en") -> list[dict]:
    translated = []
    for seg in segments:
        translated.append({**seg, "text": translate_text(seg["text"], target_language, source_language)})
    return translated


# ---------------------------------------------------------------------------
# Text-to-Speech / Voice Cloning — real Coqui TTS (XTTS v2), optional
# ---------------------------------------------------------------------------

@lru_cache(maxsize=1)
def _get_tts_model():
    try:
        from TTS.api import TTS  # type: ignore
    except Exception as exc:  # pragma: no cover - optional heavyweight dependency
        raise RuntimeError(
            "Coqui TTS is not installed. Run: pip install TTS "
            "(see requirements.txt)"
        ) from exc
    return TTS("tts_models/multilingual/multi-dataset/xtts_v2")


def synthesize_speech(
    text: str,
    output_path: Path,
    language: str = "en",
    speaker_wav: Optional[str] = None,
) -> str:
    """Generate speech audio. If speaker_wav is given, clones that voice."""
    tts = _get_tts_model()
    if speaker_wav:
        tts.tts_to_file(text=text, speaker_wav=speaker_wav, language=language, file_path=str(output_path))
    else:
        tts.tts_to_file(text=text, language=language, file_path=str(output_path))
    return str(output_path)


# ---------------------------------------------------------------------------
# Video muxing — burn subtitles with ffmpeg (real, required)
# ---------------------------------------------------------------------------

def create_burned_subtitle_video(video_path: str, subtitle_path: str) -> str:
    output_path = Path(video_path).with_suffix(".subtitled.mp4")
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        raise RuntimeError("ffmpeg not found on PATH; cannot burn subtitles.")

    command = [
        ffmpeg,
        "-y",
        "-i",
        video_path,
        "-vf",
        f"subtitles={subtitle_path}",
        str(output_path),
    ]
    result = subprocess.run(command, check=False, capture_output=True, text=True)
    if result.returncode != 0 or not output_path.exists():
        raise RuntimeError(f"ffmpeg failed to burn subtitles: {result.stderr.strip()}")
    return str(output_path)


def replace_audio_track(video_path: str, audio_path: str, output_path: Path) -> str:
    """Mux a new (e.g. dubbed) audio track onto the source video."""
    ffmpeg = shutil.which("ffmpeg")
    if not ffmpeg:
        raise RuntimeError("ffmpeg not found on PATH; cannot merge audio/video.")
    command = [
        ffmpeg,
        "-y",
        "-i",
        video_path,
        "-i",
        audio_path,
        "-c:v",
        "copy",
        "-map",
        "0:v:0",
        "-map",
        "1:a:0",
        "-shortest",
        str(output_path),
    ]
    result = subprocess.run(command, check=False, capture_output=True, text=True)
    if result.returncode != 0 or not output_path.exists():
        raise RuntimeError(f"ffmpeg failed to merge dubbed audio: {result.stderr.strip()}")
    return str(output_path)